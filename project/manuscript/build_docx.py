from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "manuscript" / "main.md"
FIGURE = ROOT / "outputs" / "reporting" / "figure_main_results.png"
OUTPUT = ROOT / "manuscript" / "rostrocaudal_eeg_validity_audit.docx"

BLUE = "2E74B5"
DARK_BLUE = "17365D"
PALE_BLUE = "EAF2F8"
PALE_GRAY = "F4F6F9"
MID_GRAY = "667085"
TEXT = "202124"
TABLE_WIDTH_DXA = 9360


def shade(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    element = properties.find(qn("w:shd"))
    if element is None:
        element = OxmlElement("w:shd")
        properties.append(element)
    element.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, bottom: int = 80, start: int = 120, end: int = 120) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for key, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = margins.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    properties = cell._tc.get_or_add_tcPr()
    width = properties.find(qn("w:tcW"))
    if width is None:
        width = OxmlElement("w:tcW")
        properties.append(width)
    width.set(qn("w:w"), str(width_dxa))
    width.set(qn("w:type"), "dxa")


def set_row_repeat(row) -> None:
    properties = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    properties.append(repeat)


def prevent_row_split(row) -> None:
    properties = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    properties.append(cant_split)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])


def add_inline(paragraph, text: str) -> None:
    text = text.replace("`", "")
    cursor = 0
    for match in re.finditer(r"\*\*(.+?)\*\*", text):
        if match.start() > cursor:
            paragraph.add_run(text[cursor : match.start()])
        run = paragraph.add_run(match.group(1))
        run.bold = True
        cursor = match.end()
    if cursor < len(text):
        paragraph.add_run(text[cursor:])


def style_run(run, size: float | None = None, color: str | None = None, bold: bool | None = None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def configure_styles(document: Document) -> None:
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    fmt = normal.paragraph_format
    fmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt.space_after = Pt(8)
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = 1.333

    heading_specs = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, DARK_BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 10, 4),
    }
    for name, (size, color, before, after) in heading_specs.items():
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = styles["Caption"]
    caption.font.name = "Calibri"
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    caption.font.size = Pt(9)
    caption.font.color.rgb = RGBColor.from_string(TEXT)
    caption.paragraph_format.space_before = Pt(6)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.keep_with_next = True

    foot = styles["Footer"]
    foot.font.name = "Calibri"
    foot.font.size = Pt(9)
    foot.font.color.rgb = RGBColor.from_string(MID_GRAY)


def add_cover(document: Document, source: str) -> None:
    lines = source.splitlines()
    title = lines[0].removeprefix("# ")
    subtitle = lines[2].removeprefix("## ")
    author = re.sub(r"\*\*", "", lines[4])
    affiliation = lines[6]
    correspondence = lines[8]
    status = re.sub(r"\*\*Draft status:\*\*\s*", "", lines[10])

    masthead = document.add_table(rows=1, cols=1)
    masthead.alignment = WD_TABLE_ALIGNMENT.CENTER
    masthead.autofit = False
    set_row_repeat(masthead.rows[0])
    set_cell_width(masthead.cell(0, 0), TABLE_WIDTH_DXA)
    shade(masthead.cell(0, 0), BLUE)
    set_cell_margins(masthead.cell(0, 0), top=110, bottom=110, start=180, end=180)
    p = masthead.cell(0, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("RESEARCH MANUSCRIPT  ·  COMPLETE ANALYSIS DRAFT")
    style_run(r, 9, "FFFFFF", True)

    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(38)
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run(title)
    style_run(r, 24, DARK_BLUE, True)

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(30)
    r = p.add_run(subtitle)
    style_run(r, 14, MID_GRAY, False)

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(author)
    style_run(r, 13, TEXT, True)
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(affiliation)
    style_run(r, 11, TEXT, False)
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(24)
    r = p.add_run(correspondence)
    style_run(r, 10, MID_GRAY, False)

    card = document.add_table(rows=1, cols=1)
    card.alignment = WD_TABLE_ALIGNMENT.CENTER
    card.autofit = False
    set_row_repeat(card.rows[0])
    set_cell_width(card.cell(0, 0), TABLE_WIDTH_DXA)
    shade(card.cell(0, 0), PALE_BLUE)
    set_cell_margins(card.cell(0, 0), top=160, bottom=160, start=220, end=220)
    p = card.cell(0, 0).paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("STATUS")
    style_run(r, 8.5, BLUE, True)
    p = card.cell(0, 0).add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(status)
    style_run(r, 9.5, TEXT, False)

    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(34)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Prepared 12 July 2026")
    style_run(r, 9.5, MID_GRAY, True)
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("88 participants  ·  5 frozen hypotheses  ·  50 outer validation folds")
    style_run(r, 9.5, MID_GRAY, False)

    document.add_page_break()


def add_markdown_table(document: Document, rows: list[list[str]], index: int) -> None:
    if index == 1:
        widths = [1800, 700, 2050, 1700, 3110]
    elif index == 2:
        widths = [1700, 2600, 1100, 1460, 1100, 1400]
    else:
        widths = [2200, 2900, 1100, 1100, 2060]
    if len(widths) != len(rows[0]):
        widths = [TABLE_WIDTH_DXA // len(rows[0])] * len(rows[0])
        widths[-1] += TABLE_WIDTH_DXA - sum(widths)

    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    table_pr = table._tbl.tblPr
    table_width = table_pr.find(qn("w:tblW"))
    if table_width is None:
        table_width = OxmlElement("w:tblW")
        table_pr.append(table_width)
    table_width.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    table_width.set(qn("w:type"), "dxa")

    for row_index, row_values in enumerate(rows):
        row = table.rows[row_index]
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        prevent_row_split(row)
        if row_index == 0:
            set_row_repeat(row)
        for column_index, value in enumerate(row_values):
            cell = row.cells[column_index]
            set_cell_width(cell, widths[column_index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index == 0:
                shade(cell, BLUE)
            elif row_index % 2 == 0:
                shade(cell, PALE_GRAY)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.LEFT if column_index == 0 else WD_ALIGN_PARAGRAPH.CENTER
            )
            run = paragraph.add_run(value.strip())
            style_run(run, 8.4, "FFFFFF" if row_index == 0 else TEXT, row_index == 0)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_figure(document: Document) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run()
    picture = run.add_picture(str(FIGURE), width=Inches(6.5))
    doc_pr = picture._inline.docPr
    doc_pr.set("descr", "Four-panel figure of complexity, state, and prediction results")
    p = document.add_paragraph(style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_inline(
        p,
        "**Figure 1. Primary results.** Direct reproduction, surrogate-normalized HFD, "
        "paired recording-state means, and participant-level cross-validated model AUCs. "
        "Full legend appears after the Conclusions.",
    )


def add_body(document: Document, source: str) -> None:
    body = source[source.index("## Abstract") :]
    lines = body.splitlines()
    table_index = 0
    figure_added = False
    in_references = False
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue
        if line.startswith("## "):
            heading = line[3:]
            if heading == "References":
                document.add_page_break()
                in_references = True
            document.add_heading(heading, level=1)
            i += 1
            continue
        if line.startswith("### "):
            heading = line[4:]
            if heading == "Multiplicity-adjusted confirmatory summary" and not figure_added:
                add_figure(document)
                figure_added = True
            document.add_heading(heading, level=2)
            i += 1
            continue
        if line.startswith("| "):
            raw_rows: list[list[str]] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                pieces = [piece.strip() for piece in lines[i].strip().strip("|").split("|")]
                if not all(re.fullmatch(r":?-{3,}:?", piece) for piece in pieces):
                    raw_rows.append(pieces)
                i += 1
            table_index += 1
            add_markdown_table(document, raw_rows, table_index)
            continue
        if line.startswith("**Table ") or line.startswith("**Figure "):
            p = document.add_paragraph(style="Caption")
            add_inline(p, line)
            i += 1
            continue
        if in_references and re.match(r"^\d+\.\s", line):
            p = document.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.first_line_indent = Inches(-0.25)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.0
            add_inline(p, line)
            for run in p.runs:
                style_run(run, 9.2, TEXT, run.bold)
            i += 1
            continue

        paragraph_lines = [line]
        i += 1
        while i < len(lines):
            nxt = lines[i].strip()
            if not nxt:
                i += 1
                break
            if nxt.startswith(("## ", "### ", "| ", "**Table ", "**Figure ")):
                break
            if re.match(r"^\d+\.\s", nxt) and line.startswith(tuple(str(n) + "." for n in range(1, 15))):
                break
            paragraph_lines.append(nxt)
            i += 1
        text = " ".join(paragraph_lines)
        p = document.add_paragraph()
        if text.startswith("**Keywords:**"):
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_inline(p, text)


def build() -> Path:
    source = SOURCE.read_text(encoding="utf-8")
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    configure_styles(document)
    core = document.core_properties
    core.title = "Testing rostrocaudal EEG complexity as a state-invariant marker of dementia subtype"
    core.subject = "EEG dementia biomarker validity audit"
    core.author = "Reuben Noronha"
    core.keywords = "EEG, Alzheimer disease, frontotemporal dementia, complexity, EEGNet"
    core.comments = "Generated from frozen pipeline outputs; institutional review required before submission."

    header = section.header
    header_table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    header_table.autofit = False
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_row_repeat(header_table.rows[0])
    for cell, width in zip(header_table.rows[0].cells, (6800, 2560), strict=True):
        set_cell_width(cell, width)
        set_cell_margins(cell, top=0, bottom=30, start=0, end=0)
    left = header_table.cell(0, 0).paragraphs[0]
    left.paragraph_format.space_after = Pt(0)
    r = left.add_run("ROSTROCAUDAL EEG COMPLEXITY VALIDITY AUDIT")
    style_run(r, 8, MID_GRAY, True)
    right = header_table.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right.paragraph_format.space_after = Pt(0)
    r = right.add_run("NORONHA")
    style_run(r, 8, MID_GRAY, True)

    footer = section.footer.paragraphs[0]
    add_page_number(footer)

    add_cover(document, source)
    add_body(document, source)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
