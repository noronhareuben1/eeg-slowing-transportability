from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "ed_freedman_pi_brief.md"
OUTPUT = ROOT / "manuscript" / "PI_Brief_Ed_Freedman_AD_FTD_EEG_Study.docx"
FIGURE = ROOT / "outputs" / "amendment_v1_2" / "figure_paired_response_results.png"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
TEXT = "20262E"
MUTED = "5F6B76"
LIGHT_GRAY = "F2F4F7"
PALE_BLUE = "E8EEF5"
WHITE = "FFFFFF"
TABLE_WIDTH_DXA = 9360


def set_run(run, *, size: float | None = None, color: str | None = None, bold=None, italic=None):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_inline(paragraph, text: str, *, size: float | None = None, color: str = TEXT) -> None:
    text = text.replace("`", "")
    cursor = 0
    for match in re.finditer(r"\*\*(.+?)\*\*", text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor : match.start()])
            set_run(run, size=size, color=color)
        run = paragraph.add_run(match.group(1))
        set_run(run, size=size, color=color, bold=True)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_run(run, size=size, color=color)


def shade(element, fill: str) -> None:
    properties = (
        element.get_or_add_tcPr()
        if hasattr(element, "get_or_add_tcPr")
        else element.get_or_add_pPr()
    )
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, *, top=80, bottom=80, start=120, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        node = OxmlElement("w:gridCol")
        node.set(qn("w:w"), str(width))
        grid.append(node)
    for row in table.rows:
        for cell, width in zip(row.cells, widths, strict=True):
            set_cell_width(cell, width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:cantSplit")
    tr_pr.append(node)


def repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run(run, size=9, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])


def create_numbering(document: Document) -> int:
    numbering = document.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), "decimal")
    level.append(fmt)
    text = OxmlElement("w:lvlText")
    text.set(qn("w:val"), "%1.")
    level.append(text)
    alignment = OxmlElement("w:lvlJc")
    alignment.set(qn("w:val"), "left")
    level.append(alignment)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    p_pr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "720")
    indent.set(qn("w:hanging"), "360")
    p_pr.append(indent)
    level.append(p_pr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    reference = OxmlElement("w:abstractNumId")
    reference.set(qn("w:val"), str(abstract_id))
    num.append(reference)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    number = OxmlElement("w:numId")
    number.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, number])


def configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.1

    specifications = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in specifications.items():
        style = document.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = document.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.tab_stops.add_tab_stop(Inches(0.5))
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        style.paragraph_format.line_spacing = 1.167

    caption = document.styles["Caption"]
    caption.font.name = "Calibri"
    caption.font.size = Pt(9)
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.keep_with_next = True

    if "Brief Quote" not in document.styles:
        quote = document.styles.add_style("Brief Quote", WD_STYLE_TYPE.PARAGRAPH)
    else:
        quote = document.styles["Brief Quote"]
    quote.font.name = "Calibri"
    quote.font.size = Pt(10.5)
    quote.font.italic = True
    quote.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    quote.paragraph_format.left_indent = Inches(0.25)
    quote.paragraph_format.right_indent = Inches(0.15)
    quote.paragraph_format.space_before = Pt(6)
    quote.paragraph_format.space_after = Pt(8)
    quote.paragraph_format.line_spacing = 1.15


def add_masthead(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(14)
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run("RESEARCH DECISION BRIEF")
    set_run(run, size=9, color=BLUE, bold=True)

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(5)
    run = paragraph.add_run("A Paired Resting-and-Photic EEG Study for AD/FTD/CN Discrimination")
    set_run(run, size=23, color=DARK_BLUE, bold=True)

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(15)
    run = paragraph.add_run("Presentation and decision brief for Ed G. Freedman")
    set_run(run, size=14, color=MUTED)

    for label, value in (
        ("Prepared by", "Reuben Noronha"),
        ("Date", "14 July 2026"),
        ("Status", "Exploratory analysis completed; confirmatory approval requested"),
    ):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        run = paragraph.add_run(f"{label}: ")
        set_run(run, size=10.5, color=TEXT, bold=True)
        run = paragraph.add_run(value)
        set_run(run, size=10.5, color=TEXT)

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(12)


def add_callout(document: Document, text: str, *, label: str | None = None) -> None:
    table = document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [TABLE_WIDTH_DXA])
    cell = table.cell(0, 0)
    shade(cell._tc, PALE_BLUE)
    set_cell_margins(cell, top=150, bottom=150, start=220, end=220)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    if label:
        run = paragraph.add_run(label.upper() + "\n")
        set_run(run, size=8.5, color=BLUE, bold=True)
    add_inline(paragraph, text, size=10.5, color=TEXT)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_markdown_table(document: Document, rows: list[list[str]]) -> None:
    column_count = len(rows[0])
    if column_count == 6:
        widths = [2200, 1432, 1432, 1432, 1432, 1432]
    else:
        base = TABLE_WIDTH_DXA // column_count
        widths = [base] * column_count
        widths[-1] += TABLE_WIDTH_DXA - sum(widths)
    table = document.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for row_index, values in enumerate(rows):
        row = table.rows[row_index]
        prevent_row_split(row)
        if row_index == 0:
            repeat_header(row)
        for column_index, value in enumerate(values):
            cell = row.cells[column_index]
            shade(cell._tc, LIGHT_GRAY if row_index == 0 else WHITE)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.LEFT if column_index == 0 else WD_ALIGN_PARAGRAPH.CENTER
            )
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            add_inline(paragraph, value.strip(), size=8.6, color=TEXT)
            if row_index == 0:
                for run in paragraph.runs:
                    run.bold = True
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_figure(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run()
    picture = run.add_picture(str(FIGURE), width=Inches(6.45))
    picture._inline.docPr.set(
        "descr",
        "Overall and class-specific ROC-AUC for resting, paired direct, "
        "two-stage, and hybrid models",
    )
    caption = document.add_paragraph(style="Caption")
    caption.alignment = WD_ALIGN_PARAGRAPH.LEFT
    add_inline(
        caption,
        "**Figure 1. Repeated participant-level validation.** Error bars show 95% "
        "participant bootstrap intervals for macro AUC; class-specific bars show "
        "one-versus-rest AUC.",
        size=9,
        color=MUTED,
    )


def _special_start(line: str) -> bool:
    stripped = line.strip()
    return bool(
        stripped.startswith(("## ", "### ", "| ", ">", "- "))
        or re.match(r"^\d+\.\s", stripped)
    )


def add_body(document: Document, source: str) -> None:
    lines = source.splitlines()
    start = next(i for i, line in enumerate(lines) if line.startswith("## The decision"))
    lines = lines[start:]
    i = 0
    result_table_seen = False
    active_num_id: int | None = None
    while i < len(lines):
        raw = lines[i]
        line = raw.strip()
        if not line:
            i += 1
            continue
        is_numbered = bool(re.match(r"^\d+\.\s", line))
        if not is_numbered:
            active_num_id = None
        if line.startswith("## "):
            heading = line[3:]
            if heading == "Likely questions and concise answers":
                document.add_page_break()
            document.add_heading(heading, level=1)
            i += 1
            continue
        if line.startswith("### "):
            heading = line[4:]
            if heading == "Locked model before test labels are opened":
                document.add_page_break()
            document.add_heading(heading, level=2)
            i += 1
            continue
        if line.startswith("| "):
            rows: list[list[str]] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                values = [value.strip() for value in lines[i].strip().strip("|").split("|")]
                if not all(re.fullmatch(r":?-{3,}:?", value) for value in values):
                    rows.append(values)
                i += 1
            add_markdown_table(document, rows)
            if not result_table_seen:
                add_figure(document)
                result_table_seen = True
            continue
        if line.startswith(">"):
            quote_lines = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                quote_lines.append(lines[i].strip().removeprefix("> "))
                i += 1
            paragraph = document.add_paragraph(style="Brief Quote")
            add_inline(paragraph, " ".join(quote_lines), size=10.5, color=DARK_BLUE)
            shade(paragraph._p, PALE_BLUE)
            continue
        if line.startswith("- "):
            text = line[2:]
            i += 1
            while i < len(lines) and lines[i].strip() and not _special_start(lines[i]):
                text += " " + lines[i].strip()
                i += 1
            paragraph = document.add_paragraph(style="List Bullet")
            add_inline(paragraph, text)
            continue
        if re.match(r"^\d+\.\s", line):
            text = re.sub(r"^\d+\.\s", "", line)
            i += 1
            while i < len(lines) and lines[i].strip() and not _special_start(lines[i]):
                text += " " + lines[i].strip()
                i += 1
            paragraph = document.add_paragraph(style="List Number")
            if active_num_id is None:
                active_num_id = create_numbering(document)
            apply_numbering(paragraph, active_num_id)
            add_inline(paragraph, text)
            continue

        paragraph_lines = [line.removesuffix("  ")]
        hard_break = raw.endswith("  ")
        i += 1
        while not hard_break and i < len(lines):
            next_raw = lines[i]
            next_line = next_raw.strip()
            if not next_line or _special_start(next_raw):
                break
            paragraph_lines.append(next_line.removesuffix("  "))
            hard_break = next_raw.endswith("  ")
            i += 1
        text = " ".join(paragraph_lines)
        if text.startswith("I am asking for approval to move this"):
            add_callout(document, text, label="Decision requested")
        else:
            paragraph = document.add_paragraph()
            add_inline(paragraph, text)


def build() -> Path:
    source = SOURCE.read_text(encoding="utf-8")
    document = Document()
    section = document.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(document)

    core = document.core_properties
    core.title = "A Paired Resting-and-Photic EEG Study for AD/FTD/CN Discrimination"
    core.subject = "PI decision and presentation brief"
    core.author = "Reuben Noronha"
    core.keywords = "EEG, Alzheimer disease, frontotemporal dementia, photic stimulation"
    core.comments = "Exploratory research brief; not a clinical diagnostic document."

    header = section.header
    table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    set_table_geometry(table, [6000, 3360])
    left = table.cell(0, 0).paragraphs[0]
    left.paragraph_format.space_after = Pt(0)
    run = left.add_run("AD/FTD/CN EEG CONFIRMATORY STUDY")
    set_run(run, size=8, color=MUTED, bold=True)
    right = table.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right.paragraph_format.space_after = Pt(0)
    run = right.add_run("PI DECISION BRIEF")
    set_run(run, size=8, color=MUTED, bold=True)
    add_page_number(section.footer.paragraphs[0])

    add_masthead(document)
    add_body(document, source)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
